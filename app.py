import os
import sys

import json
import logging
import warnings
import pandas as pd
import streamlit as st
from datetime import datetime, timezone
import re
import glob

# Ensure project root is on PYTHONPATH when running via Streamlit.
PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), os.pardir))
if PROJECT_ROOT not in sys.path:
    sys.path.insert(0, PROJECT_ROOT)

from ai_recommendation import build_recommendations
from energy_metrics import calculate_co2_tonnes, calculate_dcie, calculate_pue
from simulation import simulate_actions

logging.getLogger("pypdf").setLevel(logging.ERROR)
warnings.filterwarnings("ignore", message="Ignoring wrong pointing object*")

def ensure_dir(path: str) -> None:
    if not os.path.isdir(path):
        os.makedirs(path, exist_ok=True)


def safe_filename(name: str) -> str:
    base = os.path.basename(name)
    cleaned = "".join(ch for ch in base if ch.isalnum() or ch in {".", "_", "-"}).strip("._-")
    if not cleaned:
        cleaned = "upload"
    return cleaned


def save_uploaded_file(uploaded_file, target_dir: str) -> str:
    ensure_dir(target_dir)
    ts = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    filename = f"{ts}_{safe_filename(uploaded_file.name)}"
    target_path = os.path.join(target_dir, filename)
    with open(target_path, "wb") as handle:
        handle.write(uploaded_file.getbuffer())
    return target_path


def append_manifest(manifest_path: str, entry: dict) -> None:
    ensure_dir(os.path.dirname(manifest_path))
    data = []
    if os.path.exists(manifest_path):
        try:
            with open(manifest_path, "r", encoding="utf-8") as handle:
                data = json.load(handle)
        except Exception:
            data = []
    data.append(entry)
    with open(manifest_path, "w", encoding="utf-8") as handle:
        json.dump(data, handle, indent=2)


def cleanup_uploaded_docs() -> None:
    uploads_dir = os.path.join(PROJECT_ROOT, "uploaded_docs")
    if not os.path.isdir(uploads_dir):
        return
    for entry in os.listdir(uploads_dir):
        path = os.path.join(uploads_dir, entry)
        try:
            if os.path.isfile(path):
                os.remove(path)
        except Exception:
            continue
    try:
        if not os.listdir(uploads_dir):
            os.rmdir(uploads_dir)
    except Exception:
        pass


def upload_to_huggingface(file_path: str, repo_id: str, token: str) -> str:
    try:
        from huggingface_hub import HfApi
    except Exception:
        return "Missing dependency: huggingface_hub. Run: pip install -r requirements.txt"
    api = HfApi(token=token)
    filename = os.path.basename(file_path)
    api.upload_file(
        path_or_fileobj=file_path,
        path_in_repo=f"uploads/{filename}",
        repo_id=repo_id,
        repo_type="dataset",
    )
    return "OK"


def list_hf_files(repo_id: str, token: str) -> tuple[bool, list[str], str]:
    try:
        from huggingface_hub import HfApi
    except Exception:
        return False, [], "Missing dependency: huggingface_hub. Run: pip install -r requirements.txt"
    try:
        api = HfApi(token=token or None)
        files = api.list_repo_files(repo_id=repo_id, repo_type="dataset")
        return True, files, ""
    except Exception as exc:
        return False, [], str(exc)


def extract_metrics_from_text(text: str) -> dict:
    metrics = {}
    if not text:
        return metrics
    patterns = {
        "pue": r"\bPUE\b[^0-9]{0,20}([0-9]+(?:\.[0-9]+)?)",
        "dcie": r"\bDCiE\b[^0-9]{0,20}([0-9]+(?:\.[0-9]+)?)",
        "co2_tonnes": r"\bCO2\b[^0-9]{0,20}([0-9]+(?:\.[0-9]+)?)\s*(?:t|tonnes)",
        "it_energy_mwh": r"\bIT\s*energy(?:\s*consumption)?\b[^0-9]{0,80}([0-9]+(?:\.[0-9]+)?)\s*MWh",
        "total_energy_mwh": r"\bTotal\s*(?:data\s*center\s*)?energy(?:\s*consumption)?\b[^0-9]{0,80}([0-9]+(?:\.[0-9]+)?)\s*MWh",
        "cpu_utilization": r"\bCPU\b[^0-9]{0,10}([0-9]+(?:\.[0-9]+)?)\s*%?",
        "cooling_setpoint": r"\bCooling\s*Setpoint\b[^0-9]{0,10}([0-9]+(?:\.[0-9]+)?)",
        "virtualization_level": r"\bVirtualization\b[^0-9]{0,10}([0-9]+(?:\.[0-9]+)?)\s*%?",
        "carbon_factor": r"\bCarbon\s*Factor\b[^0-9]{0,20}([0-9]+(?:\.[0-9]+)?)\s*(?:kg|g)?\s*CO2\s*/\s*kWh",
        "latency_ms": r"\bLatency\b[^0-9]{0,10}([0-9]+(?:\.[0-9]+)?)\s*ms",
        "energy_wh_inference": r"\bEnergy\b[^0-9]{0,15}([0-9]+(?:\.[0-9]+)?)\s*Wh\s*/?\s*inference",
        "energy_kwh_inference": r"\bEnergy\b[^0-9]{0,15}([0-9]+(?:\.[0-9]+)?)\s*kWh\s*/?\s*inference",
        "cost_per_million": r"\bCost\b[^0-9]{0,20}([0-9]+(?:\.[0-9]+)?)\s*€\s*/?\s*1,?000,?000\s*inferences",
    }
    for key, pattern in patterns.items():
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            try:
                metrics[key] = float(match.group(1))
            except Exception:
                continue
    if "it_energy_mwh" in metrics and "total_energy_mwh" in metrics:
        it_energy = metrics["it_energy_mwh"]
        total_energy = metrics["total_energy_mwh"]
        if it_energy > 0:
            metrics["pue"] = total_energy / it_energy
            metrics["dcie"] = (it_energy / total_energy) * 100
    return metrics


def summarize_document(uploaded_file) -> str:
    name = uploaded_file.name.lower()
    if name.endswith(".csv"):
        try:
            if hasattr(uploaded_file, "seek"):
                uploaded_file.seek(0)
            df = pd.read_csv(uploaded_file)
            cols = ", ".join(df.columns[:12])
            return f"CSV columns: {cols}. Rows: {len(df)}."
        except Exception:
            return "CSV uploaded but could not be parsed."
    if name.endswith(".docx"):
        try:
            if hasattr(uploaded_file, "seek"):
                uploaded_file.seek(0)
            from docx import Document
        except Exception:
            return "DOCX uploaded. Install python-docx to extract text."
        try:
            doc = Document(uploaded_file)
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            content = " ".join(paragraphs[:20]).strip()
            if not content:
                return "DOCX uploaded but no readable text found."
            return "DOCX summary: " + content[:800]
        except Exception:
            return "DOCX uploaded but could not be parsed."
    if name.endswith(".pdf"):
        try:
            from pypdf import PdfReader
        except Exception:
            return "PDF uploaded. Install pypdf to extract text."
        try:
            if hasattr(uploaded_file, "seek"):
                uploaded_file.seek(0)
            reader = PdfReader(uploaded_file)
            pages_text = []
            for page in reader.pages[:3]:
                text = page.extract_text() or ""
                pages_text.append(text.strip())
            content = " ".join(pages_text).strip()
            if not content:
                return "PDF uploaded but no readable text found."
            return "PDF summary: " + content[:800]
        except Exception:
            return "PDF uploaded but could not be parsed."
    return "Unsupported document type."


def extract_pdf_text(uploaded_file, max_pages: int = 6) -> str:
    try:
        from pypdf import PdfReader
    except Exception:
        return ""
    try:
        if hasattr(uploaded_file, "seek"):
            uploaded_file.seek(0)
        reader = PdfReader(uploaded_file)
        pages_text = []
        for page in reader.pages[:max_pages]:
            text = page.extract_text() or ""
            pages_text.append(text.strip())
        return "\n".join(pages_text).strip()
    except Exception:
        return ""


def extract_docx_text(uploaded_file, max_paragraphs: int = 80) -> str:
    try:
        from docx import Document
    except Exception:
        return ""
    try:
        if hasattr(uploaded_file, "seek"):
            uploaded_file.seek(0)
        doc = Document(uploaded_file)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs[:max_paragraphs]).strip()
    except Exception:
        return ""


def extract_text_from_path(path: str, max_pages: int = 6, max_paragraphs: int = 80) -> str:
    lower = path.lower()
    if lower.endswith(".pdf"):
        try:
            from pypdf import PdfReader
        except Exception:
            return ""
        try:
            reader = PdfReader(path)
            pages_text = []
            for page in reader.pages[:max_pages]:
                text = page.extract_text() or ""
                pages_text.append(text.strip())
            return "\n".join(pages_text).strip()
        except Exception:
            return ""
    if lower.endswith(".docx"):
        try:
            from docx import Document
        except Exception:
            return ""
        try:
            doc = Document(path)
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            return "\n".join(paragraphs[:max_paragraphs]).strip()
        except Exception:
            return ""
    return ""


def load_local_knowledge_base(root_dir: str) -> dict:
    patterns = [
        os.path.join(root_dir, "*.pdf"),
        os.path.join(root_dir, "*.docx"),
    ]
    files = []
    for pattern in patterns:
        files.extend(glob.glob(pattern))
    summaries = []
    texts = []
    for path in files[:8]:
        text = extract_text_from_path(path)
        if text:
            name = os.path.basename(path)
            summaries.append(f"LOCAL DOC: {name} | {text[:400]}")
            texts.append(text)
    return {"summaries": summaries, "texts": texts}

def _to_float(value: str) -> float | None:
    if not value:
        return None
    cleaned = value.strip().replace(" ", "")
    if "," in cleaned and "." in cleaned:
        cleaned = cleaned.replace(",", "")
    elif cleaned.count(",") > 1:
        cleaned = cleaned.replace(",", "")
    elif "," in cleaned and "." not in cleaned:
        cleaned = cleaned.replace(",", ".")
    cleaned = re.sub(r"[^0-9.\-]", "", cleaned)
    try:
        return float(cleaned)
    except Exception:
        return None


def extract_workload_inputs(text: str) -> dict:
    if not text:
        return {}
    patterns = {
        "n_inferences": r"\bN\s*=\s*([0-9,.]+)\s*inferences",
        "gpu_power_w": r"\bP[_\s]*gpu\b[^0-9]{0,10}([0-9,.]+)\s*W",
        "edge_power_w": r"\bP[_\s]*edge\b[^0-9]{0,10}([0-9,.]+)\s*W",
        "gpu_latency_ms": r"\bL[_\s]*gpu\b[^0-9]{0,10}([0-9,.]+)\s*ms",
        "edge_latency_ms": r"\bL[_\s]*edge\b[^0-9]{0,10}([0-9,.]+)\s*ms",
        "gpu_cost_eur_per_hour": r"\b([0-9,.]+)\s*€\s*/\s*hour",
        "electricity_cost_eur_per_kwh": r"\b([0-9,.]+)\s*€\s*/\s*kWh",
        "hardware_cost_eur": r"\bHardware cost\b[^0-9]{0,20}([0-9,.]+)\s*€",
        "usage_per_day": r"\b([0-9,.]+)\s*inferences\s*/\s*day",
    }
    result = {}
    for key, pattern in patterns.items():
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            result[key] = _to_float(match.group(1))
    return result


def compute_workload_audit(inputs: dict) -> dict:
    n = inputs.get("n_inferences")
    gpu_p = inputs.get("gpu_power_w")
    edge_p = inputs.get("edge_power_w")
    gpu_l = inputs.get("gpu_latency_ms")
    edge_l = inputs.get("edge_latency_ms")
    gpu_cost_h = inputs.get("gpu_cost_eur_per_hour")
    elec_cost = inputs.get("electricity_cost_eur_per_kwh")
    hardware_cost = inputs.get("hardware_cost_eur")
    usage_day = inputs.get("usage_per_day")
    missing = [k for k, v in {
        "N": n,
        "P_gpu": gpu_p,
        "P_edge": edge_p,
        "L_gpu": gpu_l,
        "L_edge": edge_l,
    }.items() if v is None]
    if missing:
        return {"error": f"Missing core values: {', '.join(missing)}."}

    n = float(n)
    gpu_time_h = (n * (gpu_l / 1000.0)) / 3600.0
    edge_time_h = (n * (edge_l / 1000.0)) / 3600.0

    gpu_energy_wh = gpu_p * gpu_time_h
    edge_energy_wh = edge_p * edge_time_h
    gpu_energy_kwh = gpu_energy_wh / 1000.0
    edge_energy_kwh = edge_energy_wh / 1000.0

    gpu_energy_per_inf_wh = gpu_p * (gpu_l / 1000.0) / 3600.0
    edge_energy_per_inf_wh = edge_p * (edge_l / 1000.0) / 3600.0

    gpu_cost = gpu_cost_h * gpu_time_h if gpu_cost_h else None
    edge_cost = None
    if elec_cost:
        edge_cost = edge_energy_kwh * elec_cost
    if hardware_cost and usage_day:
        total_inf_4y = usage_day * 365 * 4
        if total_inf_4y > 0:
            edge_cost = (edge_cost or 0) + (hardware_cost / total_inf_4y) * n

    return {
        "gpu_time_h": gpu_time_h,
        "edge_time_h": edge_time_h,
        "gpu_energy_wh": gpu_energy_wh,
        "edge_energy_wh": edge_energy_wh,
        "gpu_energy_kwh": gpu_energy_kwh,
        "edge_energy_kwh": edge_energy_kwh,
        "gpu_energy_per_inf_wh": gpu_energy_per_inf_wh,
        "edge_energy_per_inf_wh": edge_energy_per_inf_wh,
        "gpu_cost": gpu_cost,
        "edge_cost": edge_cost,
    }


def summarize_workload_decision(inputs: dict, results: dict) -> str:
    lines = []
    if results["gpu_time_h"] < results["edge_time_h"]:
        lines.append("Fastest: GPU")
    else:
        lines.append("Fastest: Edge")

    if results["gpu_energy_per_inf_wh"] < results["edge_energy_per_inf_wh"]:
        lines.append("Lowest energy per inference: GPU")
    else:
        lines.append("Lowest energy per inference: Edge")

    if results["gpu_cost"] is not None and results["edge_cost"] is not None:
        if results["gpu_cost"] < results["edge_cost"]:
            lines.append("Lowest cost at scale: GPU")
        else:
            lines.append("Lowest cost at scale: Edge")
    else:
        lines.append("Cost comparison: n/a (missing cost inputs)")

    n = inputs.get("n_inferences")
    if n:
        doubled_inputs = dict(inputs)
        doubled_inputs["n_inferences"] = n * 2
        doubled = compute_workload_audit(doubled_inputs)
        if "error" not in doubled:
            if (results["gpu_cost"] is not None and results["edge_cost"] is not None and
                doubled["gpu_cost"] is not None and doubled["edge_cost"] is not None):
                change = "no"
                if (results["gpu_cost"] < results["edge_cost"]) != (doubled["gpu_cost"] < doubled["edge_cost"]):
                    change = "yes"
                lines.append(f"Does ranking change if N doubles? {change}")
    return "\n".join(f"- {line}" for line in lines)


def workload_decision_flags(results: dict) -> dict:
    fastest = "GPU" if results["gpu_time_h"] < results["edge_time_h"] else "Edge"
    lowest_energy = "GPU" if results["gpu_energy_per_inf_wh"] < results["edge_energy_per_inf_wh"] else "Edge"
    if results["gpu_cost"] is not None and results["edge_cost"] is not None:
        lowest_cost = "GPU" if results["gpu_cost"] < results["edge_cost"] else "Edge"
    else:
        lowest_cost = "n/a"
    return {
        "fastest": fastest,
        "lowest_energy": lowest_energy,
        "lowest_cost": lowest_cost,
    }


def workload_recommendation(inputs: dict, results: dict) -> str:
    fastest = "GPU" if results["gpu_time_h"] < results["edge_time_h"] else "Edge"
    lowest_energy = "GPU" if results["gpu_energy_per_inf_wh"] < results["edge_energy_per_inf_wh"] else "Edge"
    if results["gpu_cost"] is not None and results["edge_cost"] is not None:
        lowest_cost = "GPU" if results["gpu_cost"] < results["edge_cost"] else "Edge"
    else:
        lowest_cost = "n/a"
    return (
        f"- Choose {fastest} if latency/throughput is the top priority.\n"
        f"- Choose {lowest_energy} if energy per inference is the top priority.\n"
        f"- Choose {lowest_cost} if total cost at scale is the top priority."
    )


def business_goal_recommendation(goal: str, results: dict) -> str:
    flags = workload_decision_flags(results)
    if goal == "Minimum cost":
        return f"Minimum cost → Recommendation: {flags['lowest_cost']}"
    if goal == "Minimum energy per inference":
        return f"Minimum energy per inference → Recommendation: {flags['lowest_energy']}"
    if goal == "Minimum latency":
        return f"Minimum latency → Recommendation: {flags['fastest']}"
    return "Not specified"


st.set_page_config(page_title="GreenDC Audit Platform", layout="wide")
cleanup_uploaded_docs()

if "local_docs_loaded" not in st.session_state:
    local_kb = load_local_knowledge_base(PROJECT_ROOT)
    st.session_state["local_doc_texts"] = local_kb["texts"]
    st.session_state["local_docs_loaded"] = True


def find_local_excerpt(question: str, texts: list[str]) -> str | None:
    if not question or not texts:
        return None
    tokens = [t for t in re.split(r"\W+", question.lower()) if len(t) > 4]
    for text in texts:
        lower = text.lower()
        for token in tokens[:8]:
            idx = lower.find(token)
            if idx != -1:
                start = max(0, idx - 140)
                end = min(len(text), idx + 220)
                excerpt = text[start:end].strip()
                if excerpt:
                    return excerpt.replace("\n", " ")
    return None



st.markdown(
    """
    <style>
    html, body, #root, [data-testid="stAppViewContainer"], .block-container, main {
        background: #0b111f !important;
    }
    [data-testid="stSpinner"] {
        background: transparent !important;
        color: #e5ecff !important;
    }
    [data-testid="stSpinner"] * {
        color: #e5ecff !important;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

with st.sidebar:
    st.markdown("<div class='sidebar-title'>CONTROL PANEL</div>", unsafe_allow_html=True)
    if "compact_sidebar" not in st.session_state:
        st.session_state.compact_sidebar = False
    dark_mode = True
    st.caption("Theme: Dark (locked)")
    compact_sidebar = st.toggle("Compact sidebar", value=st.session_state.compact_sidebar)
    if st.button("Toggle sidebar width"):
        st.session_state.compact_sidebar = not st.session_state.compact_sidebar
        compact_sidebar = st.session_state.compact_sidebar
    st.markdown(
        """
        <div class="icon-nav">
            <a href="?page=landing" title="Landing">
                <svg width="16" height="16" viewBox="0 0 24 24" fill="none"
                 xmlns="http://www.w3.org/2000/svg"><path d="M4 10h16v10H4z" fill="#7ea6ff"/>
                 <path d="M12 4l8 6H4l8-6z" fill="#bcd0ff"/></svg>
            </a>
            <a href="?page=dashboard" title="Dashboard">
                <svg width="16" height="16" viewBox="0 0 24 24" fill="none"
                 xmlns="http://www.w3.org/2000/svg"><path d="M5 19h14v2H5z" fill="#7ea6ff"/>
                 <path d="M6 17V9h3v8H6zm5 0V5h3v12h-3zm5 0v-6h3v6h-3z" fill="#bcd0ff"/></svg>
            </a>
            <a href="?page=about" title="About">
                <svg width="16" height="16" viewBox="0 0 24 24" fill="none"
                 xmlns="http://www.w3.org/2000/svg"><circle cx="12" cy="12" r="9" stroke="#7ea6ff" stroke-width="2"/>
                 <path d="M12 8h.01M11 11h2v5h-2z" fill="#bcd0ff"/></svg>
            </a>
        </div>
        <div class="menu-item">
            <span>Sections</span>
            <span class="menu-badge">LIVE</span>
        </div>
        <div class="nav-list">
            <a href="#metrics">Metrics</a>
            <a href="#recommendations">Recommendations</a>
            <a href="#simulation">Simulation</a>
            <a href="#about">About</a>
        </div>
        """,
        unsafe_allow_html=True,
    )
    if compact_sidebar:
        st.caption("Expand sidebar to edit inputs.")
    st.markdown("<div class='menu-item'><span>GreenDC Audit AI</span><span class='menu-badge'>OFFLINE</span></div>", unsafe_allow_html=True)
    st.markdown('<div class="nav-list"><a href="#assistant">Open GreenDC Audit AI</a></div>', unsafe_allow_html=True)
    simulate_web = st.toggle("Simulate Web Search (offline)", value=False)
    if "assistant_visible" not in st.session_state:
        st.session_state.assistant_visible = True
    if st.button("Show Knowledge Assistant"):
        st.session_state.assistant_visible = True
        if hasattr(st, "query_params"):
            st.query_params["page"] = "dashboard"
        else:
            st.experimental_set_query_params(page="dashboard")
    with st.expander("Document Uploads (Audit)", expanded=False):
        st.caption("Upload PDF/CSV for audit. Files are stored locally and synced to HF if configured.")
        uploaded_docs = st.file_uploader(
            "Upload audit documents",
            type=["pdf", "csv", "docx"],
            accept_multiple_files=True,
        )
        hf_repo = "Gkop/Green-dc-audit-platform"
        hf_token = os.getenv("HF_TOKEN", "")
        if uploaded_docs:
            saved_paths = []
            doc_summaries = []
            doc_status = []
            merged_metrics = {}
            doc_texts = []
            upload_status = []
            for item in uploaded_docs:
                saved_paths.append(save_uploaded_file(item, os.path.join(PROJECT_ROOT, "uploaded_docs")))
                append_manifest(
                    os.path.join(PROJECT_ROOT, "uploaded_docs", "manifest.json"),
                    {"name": item.name, "saved_at": datetime.now(timezone.utc).isoformat()},
                )
                summary = summarize_document(item)
                doc_summaries.append(summary)
                if summary.startswith("PDF summary:"):
                    full_text = extract_pdf_text(item)
                    if full_text:
                        metrics = extract_metrics_from_text(full_text)
                        merged_metrics.update(metrics)
                        doc_texts.append(full_text)
                    doc_status.append("PDF fully readable")
                elif summary.startswith("DOCX summary:"):
                    full_text = extract_docx_text(item)
                    if full_text:
                        metrics = extract_metrics_from_text(full_text)
                        merged_metrics.update(metrics)
                        doc_texts.append(full_text)
                    doc_status.append("DOCX fully readable")
                elif summary.startswith("PDF uploaded but no readable text"):
                    doc_status.append("PDF partially readable")
                elif summary.startswith("PDF uploaded but could not be parsed"):
                    doc_status.append("PDF partially readable")
                elif summary.startswith("DOCX uploaded but no readable text"):
                    doc_status.append("DOCX partially readable")
                elif summary.startswith("DOCX uploaded but could not be parsed"):
                    doc_status.append("DOCX partially readable")
            st.success(f"Saved {len(saved_paths)} document(s) to local history.")
            if hf_token:
                synced = 0
                failed = 0
                for path in saved_paths:
                    result = upload_to_huggingface(path, hf_repo, hf_token)
                    if result == "OK":
                        synced += 1
                        upload_status.append(f"{os.path.basename(path)} • Synced to HF")
                        try:
                            os.remove(path)
                        except Exception:
                            upload_status.append(f"{os.path.basename(path)} • Local delete failed")
                    else:
                        failed += 1
                        upload_status.append(f"{os.path.basename(path)} • Not synced")
                if synced:
                    st.success(f"Synced {synced}/{len(saved_paths)} file(s) to Hugging Face.")
                    st.caption("Local files deleted after successful sync.")
                if failed:
                    st.warning(f"{failed} file(s) were not synced to Hugging Face.")
            else:
                st.error("Hugging Face token missing. Uploads were not synced and were removed locally.")
                for path in saved_paths:
                    upload_status.append(f"{os.path.basename(path)} • Not synced (no HF token)")
                    try:
                        os.remove(path)
                    except Exception:
                        upload_status.append(f"{os.path.basename(path)} • Local delete failed")
            st.session_state["doc_summaries"] = doc_summaries
            st.session_state["doc_status"] = doc_status
            st.session_state["doc_metrics"] = merged_metrics
            st.session_state["doc_texts"] = doc_texts
            st.session_state["upload_status"] = upload_status
        st.caption("No web scraping. Uses curated datasets and uploads only.")
        if st.session_state.get("doc_summaries"):
            st.markdown("<div class='menu-item'><span>Documents analyzed</span><span class='menu-badge'>OK</span></div>", unsafe_allow_html=True)
            for summary in st.session_state["doc_summaries"][:3]:
                st.caption(summary)
        if st.session_state.get("doc_status"):
            st.markdown("<div class='menu-item'><span>Doc status</span><span class='menu-badge'>INFO</span></div>", unsafe_allow_html=True)
            for status in st.session_state["doc_status"][:3]:
                st.caption(status)
        if st.session_state.get("upload_status"):
            st.markdown("<div class='menu-item'><span>Uploaded files</span><span class='menu-badge'>HF</span></div>", unsafe_allow_html=True)
            for item in st.session_state["upload_status"][:5]:
                st.caption(item)
    doc_metrics_preview = st.session_state.get("doc_metrics", {})
    applied_fields = set(doc_metrics_preview.keys()) if doc_metrics_preview else set()
    with st.expander("Energy Inputs", expanded=not compact_sidebar):
        it_label = "IT Energy (MWh/year)" + (" ✅ Applied" if "it_energy_mwh" in applied_fields else "")
        it_energy_mwh = st.number_input(
            it_label,
            min_value=0.0,
            value=float(doc_metrics_preview.get("it_energy_mwh", 780.0)),
            step=10.0,
        )
        if "it_energy_mwh" in applied_fields:
            st.markdown("<span class='applied-badge'>Applied • Data from document</span>", unsafe_allow_html=True)
        total_label = "Total Energy (MWh/year)" + (" ✅ Applied" if "total_energy_mwh" in applied_fields else "")
        total_energy_mwh = st.number_input(
            total_label,
            min_value=0.0,
            value=float(doc_metrics_preview.get("total_energy_mwh", 1300.0)),
            step=10.0,
        )
        if "total_energy_mwh" in applied_fields:
            st.markdown("<span class='applied-badge'>Applied • Data from document</span>", unsafe_allow_html=True)
        cf_label = "Carbon Factor (kg CO2/kWh)" + (" ✅ Applied" if "carbon_factor" in applied_fields else "")
        carbon_factor = st.number_input(
            cf_label,
            min_value=0.0,
            value=float(doc_metrics_preview.get("carbon_factor", 0.30)),
            step=0.01,
        )
        if "carbon_factor" in applied_fields:
            st.markdown("<span class='applied-badge'>Applied • Data from document</span>", unsafe_allow_html=True)
    with st.expander("Infrastructure Inputs", expanded=not compact_sidebar):
        servers = st.number_input("Number of Servers", min_value=0, value=320, step=10)
        cpu_label = "Average CPU Utilization (%)" + (" ✅ Applied" if "cpu_utilization" in applied_fields else "")
        cpu_utilization = st.number_input(
            cpu_label,
            min_value=0.0,
            max_value=100.0,
            value=float(doc_metrics_preview.get("cpu_utilization", 18.0)),
        )
        if "cpu_utilization" in applied_fields:
            st.markdown("<span class='applied-badge'>Applied • Data from document</span>", unsafe_allow_html=True)
        virt_label = "Virtualization Level (%)" + (" ✅ Applied" if "virtualization_level" in applied_fields else "")
        virtualization_level = st.number_input(
            virt_label,
            min_value=0.0,
            max_value=100.0,
            value=float(doc_metrics_preview.get("virtualization_level", 45.0)),
        )
        if "virtualization_level" in applied_fields:
            st.markdown("<span class='applied-badge'>Applied • Data from document</span>", unsafe_allow_html=True)
    with st.expander("Cooling & Facilities", expanded=not compact_sidebar):
        cool_label = "Cooling Setpoint (°C)" + (" ✅ Applied" if "cooling_setpoint" in applied_fields else "")
        cooling_setpoint = st.number_input(
            cool_label,
            min_value=10.0,
            max_value=30.0,
            value=float(doc_metrics_preview.get("cooling_setpoint", 19.0)),
        )
        if "cooling_setpoint" in applied_fields:
            st.markdown("<span class='applied-badge'>Applied • Data from document</span>", unsafe_allow_html=True)
        aisle_containment = st.checkbox("Hot/Cold Aisle Containment in place", value=False)

doc_metrics = st.session_state.get("doc_metrics", {})
workload_metrics = {
    "latency_ms": doc_metrics.get("latency_ms"),
    "energy_wh_inference": doc_metrics.get("energy_wh_inference"),
    "energy_kwh_inference": doc_metrics.get("energy_kwh_inference"),
    "cost_per_million": doc_metrics.get("cost_per_million"),
}
has_workload_metrics = any(value is not None for value in workload_metrics.values())
applied_params = []
applied_fields = set()
if doc_metrics:
    it_energy_mwh = doc_metrics.get("it_energy_mwh", it_energy_mwh)
    total_energy_mwh = doc_metrics.get("total_energy_mwh", total_energy_mwh)
    carbon_factor = doc_metrics.get("carbon_factor", carbon_factor)
    cpu_utilization = doc_metrics.get("cpu_utilization", cpu_utilization)
    cooling_setpoint = doc_metrics.get("cooling_setpoint", cooling_setpoint)
    virtualization_level = doc_metrics.get("virtualization_level", virtualization_level)
    if "it_energy_mwh" in doc_metrics:
        applied_params.append(f"IT energy applied from doc: {doc_metrics['it_energy_mwh']:.0f} MWh/year")
        applied_fields.add("it_energy_mwh")
    if "total_energy_mwh" in doc_metrics:
        applied_params.append(f"Total energy applied from doc: {doc_metrics['total_energy_mwh']:.0f} MWh/year")
        applied_fields.add("total_energy_mwh")
    if "carbon_factor" in doc_metrics:
        applied_params.append(f"Carbon factor applied from doc: {doc_metrics['carbon_factor']:.2f} kg CO2/kWh")
        applied_fields.add("carbon_factor")
    if "cpu_utilization" in doc_metrics:
        applied_params.append(f"CPU utilization applied from doc: {doc_metrics['cpu_utilization']:.1f}%")
        applied_fields.add("cpu_utilization")
    if "cooling_setpoint" in doc_metrics:
        applied_params.append(f"Cooling setpoint applied from doc: {doc_metrics['cooling_setpoint']:.1f} °C")
        applied_fields.add("cooling_setpoint")
    if "virtualization_level" in doc_metrics:
        applied_params.append(f"Virtualization level applied from doc: {doc_metrics['virtualization_level']:.1f}%")
        applied_fields.add("virtualization_level")

theme = "dark" if dark_mode else "light"
if dark_mode:
    bg = "radial-gradient(140% 140% at 0% 0%, #0f1d3b 0%, #0a1022 55%, #060914 100%)"
    text = "#e9edff"
    muted = "#c9d4ff"
    panel = "rgba(18, 26, 51, 0.92)"
    panel_border = "rgba(255,255,255,0.12)"
    card_bg = "linear-gradient(180deg, rgba(24, 36, 72, 0.98) 0%, rgba(12, 20, 40, 0.98) 100%)"
    hover_bg = "rgba(28, 40, 78, 0.95)"
else:
    # Soft-dark corporate palette to avoid white glare.
    bg = "radial-gradient(120% 120% at 0% 0%, #121a2c 0%, #0f1628 60%, #0b111f 100%)"
    text = "#e5ecff"
    muted = "#b6c5e6"
    panel = "rgba(21, 30, 54, 0.92)"
    panel_border = "rgba(255,255,255,0.12)"
    card_bg = "linear-gradient(180deg, rgba(28, 40, 72, 0.98) 0%, rgba(16, 24, 48, 0.98) 100%)"
    hover_bg = "rgba(32, 46, 86, 0.95)"

st.markdown(
    f"""
    <style>
    html {{
        scroll-behavior: smooth;
    }}
    .stApp {{
        background: {bg};
    }}
    body, [data-testid="stAppViewContainer"] {{
        color: {text};
    }}
    [data-testid="stHeader"], [data-testid="stToolbar"], [data-testid="stDecoration"] {{
        background: {panel} !important;
        border-bottom: 1px solid {panel_border} !important;
    }}
    [data-testid="stAppViewContainer"], .block-container {{
        background: transparent !important;
    }}
    [data-testid="stSidebar"] {{
        background: {panel} !important;
    }}
    h1, h2, h3, h4, h5, p, li, label, span, div {{
        color: {text};
    }}
    a {{ color: {muted}; }}
    [data-testid="stMetricValue"] {{ color: {text} !important; text-shadow: 0 2px 6px rgba(0,0,0,0.25); }}
    [data-testid="stMetricLabel"] {{ color: {muted} !important; }}
    input, textarea, select, button {{
        background: {panel} !important;
        color: {text} !important;
        border: 1px solid {panel_border} !important;
    }}
    input:hover, textarea:hover, select:hover, button:hover {{
        background: {hover_bg} !important;
        color: {text} !important;
    }}
    input:focus, textarea:focus, select:focus {{
        outline: none !important;
        box-shadow: 0 0 0 2px {"rgba(126, 166, 255, 0.45)" if dark_mode else "rgba(76, 130, 255, 0.35)"} !important;
    }}
    [data-baseweb="select"] * {{
        color: {text} !important;
    }}
    [data-baseweb="input"] input {{
        background: {panel} !important;
        color: {text} !important;
    }}
    [data-baseweb="select"] div {{
        background: {panel} !important;
        color: {text} !important;
    }}
    [data-baseweb="popover"] {{
        background: {panel} !important;
    }}
    [data-baseweb="popover"] > div {{
        background: {panel} !important;
        color: {text} !important;
    }}
    [data-baseweb="popover"] * {{
        color: {text} !important;
        background-color: transparent !important;
    }}
    [data-baseweb="tooltip"], [data-baseweb="modal"], [data-testid="stModal"] {{
        background: {panel} !important;
        color: {text} !important;
        border: 1px solid {panel_border} !important;
    }}
    [data-baseweb="modal"] * {{
        color: {text} !important;
        background: transparent !important;
    }}
    [role="dialog"] {{
        background: {panel} !important;
        color: {text} !important;
        border: 1px solid {panel_border} !important;
    }}
    [role="dialog"] * {{
        color: {text} !important;
        background: transparent !important;
    }}
    [data-baseweb="layer"] {{
        background: transparent !important;
    }}
    [data-baseweb="menu"] {{
        background: {panel} !important;
        border: 1px solid {panel_border} !important;
    }}
    [role="menu"] {{
        background: {panel} !important;
        color: {text} !important;
    }}
    [role="menu"] li, [role="menuitem"] {{
        background: {panel} !important;
        color: {text} !important;
    }}
    [role="menuitem"]:hover {{
        background: {hover_bg} !important;
    }}
    [data-testid="stAppMenuPopover"], [data-testid="stPopoverBody"] {{
        background: {panel} !important;
        color: {text} !important;
        border: 1px solid {panel_border} !important;
    }}
    [data-testid="stAppMenuPopover"] * {{
        color: {text} !important;
        background: transparent !important;
    }}
    [data-testid="stAppMenuPopover"] div[role="menuitem"] {{
        background: {panel} !important;
        color: {text} !important;
    }}
    [data-testid="stAppMenuPopover"] div[role="menuitem"]:hover {{
        background: {hover_bg} !important;
        color: {text} !important;
    }}
    [data-testid="stAppMenuPopover"] hr {{
        border-color: {panel_border} !important;
    }}
    [data-testid="stAppMenuPopover"] kbd {{
        background: {panel} !important;
        color: {text} !important;
        border: 1px solid {panel_border} !important;
    }}
    [data-testid="stAppMenuPopover"] a:hover {{
        background: {hover_bg} !important;
    }}
    [data-testid="stAppToolbar"] button {{
        background: {panel} !important;
        color: {text} !important;
        border: 1px solid {panel_border} !important;
    }}
    [data-testid="stAppToolbar"] svg {{
        fill: {text} !important;
        color: {text} !important;
    }}
    [data-testid="stToggle"] div {{
        background: {panel} !important;
        border: 1px solid {panel_border} !important;
    }}
    [data-testid="stRadio"] div {{
        background: transparent !important;
    }}
    [data-testid="stRadio"] label {{
        color: {text} !important;
    }}
    .section-title {{
        font-size: 18px;
        font-weight: 700;
        letter-spacing: 0.02em;
        margin: 6px 0 10px 0;
        color: {text};
        text-shadow: 0 2px 6px rgba(0,0,0,0.4);
    }}
    .section-title {{
        border-left: 4px solid #7ee6b4;
        padding-left: 10px;
    }}
    .topbar {{
        position: sticky;
        top: 0;
        z-index: 999;
        background: {panel};
        backdrop-filter: blur(10px);
        border: 1px solid rgba(126, 230, 180, 0.45);
        border-radius: 14px;
        padding: 10px 14px;
        margin: 6px 0 16px 0;
        box-shadow: 0 8px 20px rgba(0,0,0,0.35), 0 0 0 1px rgba(126, 230, 180, 0.15);
        display: flex;
        align-items: center;
        justify-content: space-between;
        gap: 12px;
    }}
    .logo {{
        display: flex;
        align-items: center;
        gap: 8px;
        font-weight: 800;
        letter-spacing: 0.02em;
        color: {text};
    }}
    .logo span {{
        color: {muted};
        font-weight: 600;
        font-size: 12px;
        margin-left: 4px;
    }}
    .nav a {{
        text-decoration: none;
        margin-right: 12px;
        padding: 6px 10px;
        border-radius: 10px;
        background: {panel};
        border: 1px solid rgba(126, 230, 180, 0.35);
        color: {text};
        font-size: 12px;
    }}
    .nav a:hover {{
        background: {hover_bg};
        color: {text};
    }}
    .dropdown {{
        position: relative;
        display: inline-block;
    }}
    .dropdown-content {{
        display: none;
        position: absolute;
        min-width: 160px;
        background: {panel};
        border: 1px solid rgba(126, 230, 180, 0.35);
        border-radius: 12px;
        box-shadow: 0 10px 20px rgba(0,0,0,0.35);
        padding: 6px;
        z-index: 1000;
    }}
    .dropdown:hover .dropdown-content {{
        display: block;
    }}
    .dropdown .dropdown-content .dropdown {{
        position: relative;
    }}
    .dropdown .dropdown-content .dropdown .dropdown-content {{
        left: 160px;
        top: 0;
    }}
    .dropdown-content a {{
        display: block;
        padding: 6px 10px;
        margin: 4px 0;
        border-radius: 8px;
        background: {panel};
        color: {text};
        text-decoration: none;
        font-size: 12px;
        font-weight: 600;
    }}
    .dropdown-content a:hover {{
        background: {hover_bg};
        color: {text};
    }}
    .subtle {{
        opacity: 0.8;
        font-size: 13px;
    }}
    .hero {{
        background: {card_bg};
        color: #f7f7f7;
        padding: 28px 32px;
        border-radius: 16px;
        box-shadow: 0 10px 30px rgba(0,0,0,0.35);
        border: 1px solid {panel_border};
        margin-bottom: 20px;
        position: relative;
        overflow: hidden;
    }}
    .hero h1 {{ margin: 0 0 6px 0; font-size: 32px; }}
    .hero p {{ margin: 0; opacity: 0.9; }}
    .hero::after {{
        content: "";
        position: absolute;
        right: -80px;
        top: -80px;
        width: 220px;
        height: 220px;
        border-radius: 50%;
        background: radial-gradient(circle, rgba(110,140,255,0.35), rgba(110,140,255,0.0) 60%);
        filter: blur(2px);
    }}
    .badge {{
        display: inline-block;
        padding: 6px 12px;
        border-radius: 999px;
        background: {panel};
        border: 1px solid {panel_border};
        color: {text};
        font-size: 12px;
        letter-spacing: 0.04em;
        margin-right: 8px;
    }}
    .badge-solid {{
        background: rgba(76, 214, 180, 0.18);
        border: 1px solid rgba(76, 214, 180, 0.45);
        color: #c9fff2;
    }}
    .glass {{
        background: {panel};
        backdrop-filter: blur(10px);
        border: 1px solid {panel_border};
        box-shadow: 0 12px 24px rgba(0,0,0,0.25);
        border-radius: 16px;
        padding: 16px 18px;
    }}
    .metric-card {{
        background: {card_bg};
        color: {text};
        padding: 16px 18px;
        border-radius: 14px;
        box-shadow: 0 12px 28px rgba(0,0,0,0.45);
        border: 1px solid {panel_border};
    }}
    .metric-card h3 {{ margin: 0 0 6px 0; font-size: 12px; text-transform: uppercase; letter-spacing: 0.1em; color: {muted}; }}
    .metric-card .value {{ font-size: 28px; font-weight: 800; color: {text}; text-shadow: 0 2px 6px rgba(0,0,0,0.25); }}
    .stSidebar > div:first-child {{
        background: {panel};
        border-right: 1px solid {panel_border};
    }}
    .stSidebar label, .stSidebar span, .stSidebar p {{
        color: {text} !important;
    }}
    .sidebar-title {{
        font-weight: 800;
        font-size: 14px;
        letter-spacing: 0.08em;
        margin-bottom: 8px;
        color: {muted};
    }}
    .stSidebar [data-testid="stExpander"] {{
        background: {panel};
        border-radius: 12px;
        border: 1px solid {panel_border};
        padding: 6px;
    }}
    .stSidebar [data-testid="stExpander"] summary {{
        background: {panel} !important;
        color: {text} !important;
        border-radius: 10px;
    }}
    .stSidebar [data-testid="stExpander"] summary:hover {{
        background: {hover_bg} !important;
    }}
    .stSidebar [data-testid="stExpander"] summary {{
        color: #e9edff;
        font-weight: 600;
    }}
    .stSidebar [role="radiogroup"] label {{
        color: {muted} !important;
        font-weight: 600;
    }}
    .stSidebar [role="radiogroup"] label[data-selected="true"] {{
        color: {text} !important;
    }}
    .icon-nav {{
        display: grid;
        grid-template-columns: repeat(3, 1fr);
        gap: 8px;
        margin-bottom: 8px;
    }}
    .icon-nav a {{
        display: inline-flex;
        justify-content: center;
        align-items: center;
        height: 34px;
        border-radius: 10px;
        background: {panel};
        border: 1px solid {panel_border};
        text-decoration: none;
    }}
    .menu-item {{
        display: flex;
        align-items: center;
        gap: 8px;
        padding: 6px 8px;
        border-radius: 10px;
        background: {"rgba(20, 30, 60, 0.6)" if dark_mode else "rgba(230, 236, 255, 0.9)"};
        border: 1px solid rgba(126, 230, 180, 0.35);
        color: {text};
        margin-bottom: 6px;
        font-size: 12px;
    }}
    .menu-badge {{
        margin-left: auto;
        font-size: 10px;
        padding: 2px 8px;
        border-radius: 999px;
        background: {"rgba(120, 210, 255, 0.15)" if dark_mode else "rgba(76, 214, 180, 0.18)"};
        border: 1px solid {"rgba(120, 210, 255, 0.35)" if dark_mode else "rgba(76, 214, 180, 0.45)"};
        color: {text};
    }}
    .nav-list a {{
        display: block;
        padding: 6px 10px;
        margin: 4px 0;
        border-radius: 10px;
        background: {panel};
        border: 1px solid rgba(126, 230, 180, 0.28);
        color: {text};
        text-decoration: none;
        font-size: 12px;
        font-weight: 600;
    }}
    .nav-list a:hover {{
        background: {hover_bg};
        color: {text};
    }}
    .stDataFrame, .stTable {{
        background: rgba(20, 30, 60, 0.35);
        border-radius: 12px;
        padding: 6px;
    }}
    .stDataFrame div[role="grid"] {{
        background: rgba(16, 24, 48, 0.9);
        color: #e9edff;
    }}
    .stDataFrame div[role="grid"] * {{
        color: #e9edff !important;
    }}
    .stTabs [data-baseweb="tab"] {{
        color: #c9d4ff;
        background: rgba(20, 30, 60, 0.5);
        border-radius: 12px;
        margin-right: 6px;
        padding: 8px 14px;
        border: 1px solid rgba(255,255,255,0.08);
    }}
    .stTabs [aria-selected="true"] {{
        background: rgba(52, 75, 140, 0.7);
        color: #ffffff;
        border: 1px solid rgba(255,255,255,0.2);
        box-shadow: 0 6px 16px rgba(0,0,0,0.35);
    }}
    .footer {{
        margin-top: 20px;
        padding: 12px 16px;
        border-radius: 12px;
        background: {panel};
        border: 1px solid {panel_border};
        font-size: 12px;
        color: {muted};
    }}
    .section {{
        background: {panel};
        border: 1px solid rgba(126, 230, 180, 0.25);
        border-radius: 16px;
        padding: 18px;
        margin-bottom: 16px;
        box-shadow: 0 10px 20px rgba(0,0,0,0.35);
    }}
    .section + .section {{
        margin-top: 12px;
    }}
    .section-title {{
        margin-top: 18px;
        margin-bottom: 12px;
    }}
    .soft-divider {{
        height: 1px;
        background: {panel_border};
        margin: 14px 0;
    }}
    .applied-badge {{
        color: #7ee6b4;
        font-weight: 700;
        font-size: 11px;
    }}
    .kpi-applied {{
        display: inline-block;
        margin-left: 6px;
        padding: 2px 6px;
        border-radius: 999px;
        font-size: 10px;
        font-weight: 700;
        color: #0b1a14;
        background: #7ee6b4;
        box-shadow: 0 0 0 1px rgba(126, 230, 180, 0.65), 0 0 10px rgba(126, 230, 180, 0.45);
    }}
    .kpi-source {{
        margin-top: 6px;
        font-size: 11px;
        color: #b6f5dc;
    }}
    .rec-card {{
        background: {card_bg};
        border: 1px solid {panel_border};
        border-radius: 14px;
        padding: 14px 16px;
        margin-bottom: 10px;
        box-shadow: 0 10px 20px rgba(0,0,0,0.35);
    }}
    .rec-title {{
        font-weight: 700;
        font-size: 14px;
        margin-bottom: 6px;
        color: #f1f4ff;
    }}
    .rec-meta {{
        display: inline-block;
        padding: 4px 10px;
        border-radius: 999px;
        background: {panel};
        border: 1px solid {panel_border};
        color: {text};
        font-size: 11px;
        margin-top: 8px;
    }}
    .svg-icon {{
        vertical-align: middle;
        margin-right: 6px;
    }}
    .action-icon {{
        vertical-align: middle;
        margin-right: 8px;
    }}
    .compact .stSidebar {{
        width: 72px !important;
    }}
    .compact .stSidebar label, .compact .stSidebar span, .compact .stSidebar p {{
        font-size: 10px;
    }}
    {".compact [data-testid='stSidebar'] { width: 72px !important; }" if compact_sidebar else ""}
    </style>
    """,
    unsafe_allow_html=True,
)

if compact_sidebar:
    st.markdown("<div class='compact'></div>", unsafe_allow_html=True)

if hasattr(st, "query_params"):
    query = st.query_params
    page_param = query.get("page", "landing")
    if isinstance(page_param, list):
        page_param = page_param[0]
else:
    query = st.experimental_get_query_params()
    page_param = query.get("page", ["landing"])[0]
page_param = page_param.lower()
if page_param not in {"landing", "dashboard", "about"}:
    page_param = "landing"
page = page_param.capitalize()


def action_icon_svg(action_title: str) -> str:
    title = action_title.lower()
    if "consolidation" in title:
        return (
            "<svg class='action-icon' width='16' height='16' viewBox='0 0 24 24' fill='none' "
            "xmlns='http://www.w3.org/2000/svg'><path d='M4 4h16v6H4z' fill='#cfe0ff'/>"
            "<path d='M6 13h12v7H6z' fill='#9fb2ff'/></svg>"
        )
    if "cooling" in title:
        return (
            "<svg class='action-icon' width='16' height='16' viewBox='0 0 24 24' fill='none' "
            "xmlns='http://www.w3.org/2000/svg'><path d='M12 3v18' stroke='#cfe0ff' stroke-width='2'/>"
            "<path d='M8 7h8M8 12h8M8 17h8' stroke='#cfe0ff' stroke-width='2'/></svg>"
        )
    if "aisle" in title:
        return (
            "<svg class='action-icon' width='16' height='16' viewBox='0 0 24 24' fill='none' "
            "xmlns='http://www.w3.org/2000/svg'><path d='M4 4h6v16H4z' fill='#cfe0ff'/>"
            "<path d='M14 4h6v16h-6z' fill='#9fb2ff'/></svg>"
        )
    if "virtualization" in title:
        return (
            "<svg class='action-icon' width='16' height='16' viewBox='0 0 24 24' fill='none' "
            "xmlns='http://www.w3.org/2000/svg'><rect x='4' y='4' width='7' height='7' fill='#cfe0ff'/>"
            "<rect x='13' y='4' width='7' height='7' fill='#9fb2ff'/>"
            "<rect x='4' y='13' width='7' height='7' fill='#9fb2ff'/></svg>"
        )
    return (
        "<svg class='action-icon' width='16' height='16' viewBox='0 0 24 24' fill='none' "
        "xmlns='http://www.w3.org/2000/svg'><circle cx='12' cy='12' r='8' stroke='#cfe0ff' stroke-width='2'/></svg>"
    )


def is_audit_question(question: str) -> bool:
    q = question.lower()
    keywords = [
        "pue",
        "dcie",
        "co2",
        "carbon",
        "energy",
        "audit",
        "data center",
        "cooling",
        "virtualization",
        "consolidation",
        "server",
        "efficiency",
        "recommendation",
        "simulation",
        "reduce",
        "optimize",
    ]
    return any(k in q for k in keywords)


def is_definition_question(question: str) -> bool:
    q = question.lower()
    return any(k in q for k in ["what is", "define", "definition", "explain"]) and (
        "green it" in q or "sustainable energy" in q
    )


def is_document_question(question: str) -> bool:
    q = question.lower()
    return any(k in q for k in ["pdf", "document", "doc", "chapter", "summary", "extracted"])


def is_long_paste(question: str) -> bool:
    return len(question.strip()) > 300


def ai_assistant_reply(question: str, context: dict) -> str:
    if not question.strip():
        return "Please enter a question related to Green IT audits or data center optimization."
    if is_definition_question(question):
        excerpt = find_local_excerpt(question, st.session_state.get("local_doc_texts", []))
        if excerpt:
            return (
                "Based on course materials:\n"
                + excerpt
                + "\n\nGreen IT reduces the environmental impact of digital systems across "
                "their lifecycle, while sustainable energy focuses on low‑emission energy sources "
                "and efficient consumption."
            )
        return (
            "Green IT focuses on reducing the environmental impact of digital systems across "
            "their entire lifecycle (design, use, and end-of-life) through sobriety, efficiency, "
            "and measurable KPIs. Sustainable energy refers to energy produced and used in ways "
            "that minimize emissions and preserve resources for the long term (renewables, "
            "efficient consumption, and responsible sourcing)."
        )
    docs = context.get("doc_summaries", [])
    if is_document_question(question) or is_long_paste(question):
        extracted = extract_metrics_from_text(question)
        workload_inputs = extract_workload_inputs(question)
        if extracted.get("it_energy_mwh") and extracted.get("total_energy_mwh"):
            it_energy = extracted["it_energy_mwh"]
            total_energy = extracted["total_energy_mwh"]
            carbon_factor = extracted.get("carbon_factor", context["carbon_factor"])
            pue = total_energy / it_energy if it_energy else context["pue"]
            dcie = (it_energy / total_energy) * 100 if total_energy else context["dcie"]
            co2 = calculate_co2_tonnes(total_energy, carbon_factor)
            intro = (
                "Detected audit metrics from the pasted document:\n"
                f"PUE {pue:.2f}, DCiE {dcie:.1f}%, CO2 {co2:.1f} t/y, "
                f"IT energy {it_energy:.0f} MWh/y, total energy {total_energy:.0f} MWh/y."
            )
            return intro
        if any(workload_inputs.get(k) is not None for k in ["n_inferences", "gpu_power_w", "edge_power_w", "gpu_latency_ms", "edge_latency_ms"]):
            return "Use the Document QA block below to analyze AI workload benchmarks."
        return "Use the Document QA block below to analyze uploaded documents."
    if not is_audit_question(question) and not context.get("doc_metrics"):
        return (
            "I can only answer questions related to Green IT audits, data center energy, "
            "and optimization within this platform."
        )

    it_energy = context["it_energy_mwh"]
    total_energy = context["total_energy_mwh"]
    carbon_factor = context["carbon_factor"]
    pue = context["pue"]
    dcie = context["dcie"]
    co2 = context["co2_tonnes"]
    cpu = context["cpu_utilization"]
    cooling = context["cooling_setpoint"]
    aisle = context["aisle_containment"]
    virt = context["virtualization_level"]
    recs = context["recommendations"]
    docs = context.get("doc_summaries", [])

    intro = (
        f"Based on your audit inputs: PUE {pue:.2f}, DCiE {dcie:.1f}%, CO2 {co2:.1f} t/y, "
        f"IT energy {it_energy:.0f} MWh/y, total energy {total_energy:.0f} MWh/y."
    )
    ops = []
    if cpu < 30:
        ops.append("consolidate servers to reduce idle energy")
    if cooling < 22:
        ops.append("raise cooling setpoint to reduce cooling load")
    if not aisle:
        ops.append("add hot/cold aisle containment to improve airflow efficiency")
    if virt < 60:
        ops.append("increase virtualization to reduce physical footprint")
    if not ops:
        ops.append("maintain current configuration and keep monitoring")

    plan = "Improvement plan: " + "; ".join(ops) + "."
    note = (
        "This assistant is a built-in, offline advisor. It does not access the internet "
        "and only uses the platform context."
    )
    return f"{intro}\n\n{plan}\n\n{note}"


def load_rules() -> dict:
    rules_path = os.path.join(PROJECT_ROOT, "knowledge_base", "rules.json")
    try:
        with open(rules_path, "r", encoding="utf-8") as handle:
            return json.load(handle)
    except FileNotFoundError:
        return {"rules": [], "standards": []}


def rule_based_plan(context: dict, rules: dict) -> str:
    cpu = context["cpu_utilization"]
    cooling_ratio = context["cooling_ratio"]
    pue = context["pue"]
    applied = []
    for rule in rules.get("rules", []):
        if rule["id"] == "CPU_LOW" and cpu < 20:
            applied.append(rule)
        if rule["id"] == "COOLING_HIGH" and cooling_ratio > 60:
            applied.append(rule)
        if rule["id"] == "PUE_HIGH" and pue > 1.6:
            applied.append(rule)
    if not applied:
        return "No rule triggered. Keep monitoring and maintain current best practices."
    lines = []
    for rule in applied:
        lines.append(
            f"- {rule['action']} (~{rule['estimated_energy_saving_percent']}%): {rule['justification']}"
        )
    return "Rule-based actions:\n" + "\n".join(lines)


 


if page == "Landing":
    st.markdown(
        """
        <div class="hero">
            <h1>🌿⚡ GreenDC Audit Platform</h1>
            <p>AI-assisted energy & carbon audit for industrial data centers — modern, 3D-inspired, and Green IT compliant.</p>
            <div style="margin-top: 12px;">
                <span class="badge">Green IT</span>
                <span class="badge">Green Coding</span>
                <span class="badge">ISO 50001 Ready</span>
                <span class="badge">-25% CO2 Target</span>
            </div>
            <div style="margin-top: 10px; color:#7ee6b4; font-weight:700;">🍃 GreenAI Systems</div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    left, right = st.columns([1.2, 1])
    with left:
        st.markdown("<div class='section-title'>Why GreenDC?</div>", unsafe_allow_html=True)
        st.markdown(
            """
            <div class="section">
                A practical platform to measure, explain, and validate energy and carbon reductions.
                Built for industrial data centers with realistic constraints.
            </div>
            """,
            unsafe_allow_html=True,
        )
        st.markdown("")
        st.markdown(
            """
            <div class="section">
                <b>Core pillars:</b>
                <ul>
                    <li>Measurable KPIs (PUE, DCiE, CO2)</li>
                    <li>Actionable recommendations</li>
                    <li>Scenario validation to reach -25%</li>
                </ul>
            </div>
            """,
            unsafe_allow_html=True,
        )
    with right:
        st.markdown("<div class='section-title'>Platform Snapshot</div>", unsafe_allow_html=True)
        st.markdown(
            """
            <div class="section">
                <b>Inputs:</b> Energy, cooling, utilization, carbon factor<br><br>
                <b>Outputs:</b> KPIs, AI recommendations, savings simulation<br><br>
                <span class="badge badge-solid">Audit-ready</span>
                <span class="badge badge-solid">Explainable</span>
                <span class="badge badge-solid">Lightweight</span>
            </div>
            """,
            unsafe_allow_html=True,
        )

if page == "Dashboard":
    st.markdown(
        """
        <div class="topbar">
            <div class="logo">
                <svg width="26" height="26" viewBox="0 0 24 24" fill="none"
                 xmlns="http://www.w3.org/2000/svg">
                  <defs>
                    <linearGradient id="leaf" x1="0" y1="0" x2="1" y2="1">
                      <stop offset="0%" stop-color="#7ee6b4"/>
                      <stop offset="100%" stop-color="#4cd6b4"/>
                    </linearGradient>
                    <linearGradient id="dc" x1="0" y1="0" x2="0" y2="1">
                      <stop offset="0%" stop-color="#cfe0ff"/>
                      <stop offset="100%" stop-color="#9fb2ff"/>
                    </linearGradient>
                    <linearGradient id="node" x1="0" y1="0" x2="1" y2="1">
                      <stop offset="0%" stop-color="#6fd3ff"/>
                      <stop offset="100%" stop-color="#3e7bff"/>
                    </linearGradient>
                  </defs>
                  <path d="M5 4h14v6H5z" fill="url(#dc)"/>
                  <path d="M4 12h16v7H4z" fill="url(#dc)"/>
                  <path d="M12 2c3.2 1 4.5 3.2 4.5 5.4-2-1.2-4.5-1.2-6.8 0 0-2.2 1.3-4.4 2.3-5.4z"
                        fill="url(#leaf)"/>
                  <circle cx="19" cy="18" r="2" fill="url(#node)"/>
                  <circle cx="5" cy="18" r="2" fill="url(#node)"/>
                  <path d="M7 18h10" stroke="#6fd3ff" stroke-width="1.5"/>
                </svg>
                🌿⚡ <span style="color:#7ee6b4;">GreenDC</span> <span>Audit Console</span>
            </div>
            <div class="nav">
                <div class="dropdown">
                    <a href="#metrics">KPIs ▾</a>
                    <div class="dropdown-content">
                        <a href="#metrics">Energy KPIs</a>
                        <a href="#recommendations">AI Actions</a>
                        <a href="#simulation">Impact</a>
                        <div class="dropdown">
                            <a href="#metrics">More ▸</a>
                            <div class="dropdown-content">
                                <a href="#metrics">Efficiency</a>
                                <a href="#recommendations">Optimization</a>
                            </div>
                        </div>
                    </div>
                </div>
                <div class="dropdown">
                    <a href="?page=about#about">Platform ▾</a>
                    <div class="dropdown-content">
                        <a href="?page=about#about">About</a>
                        <a href="#simulation">Simulation</a>
                        <div class="dropdown">
                            <a href="?page=about#team">Team ▸</a>
                            <div class="dropdown-content">
                                <a href="?page=about#team">GreenAI Systems 🌱</a>
                            </div>
                        </div>
                    </div>
                </div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    st.markdown("<div id='metrics' class='section-title'>Key Metrics</div>", unsafe_allow_html=True)
    metrics_col, recs_col = st.columns([1, 1])

    with metrics_col:
        pue = calculate_pue(it_energy_mwh, total_energy_mwh)
        dcie = calculate_dcie(it_energy_mwh, total_energy_mwh)
        co2_tonnes = calculate_co2_tonnes(total_energy_mwh, carbon_factor)
        if doc_metrics:
            if "pue" in doc_metrics:
                pue = doc_metrics["pue"]
            if "dcie" in doc_metrics:
                dcie = doc_metrics["dcie"]
            if "co2_tonnes" in doc_metrics:
                co2_tonnes = doc_metrics["co2_tonnes"]
        pue_from_doc = "pue" in doc_metrics
        dcie_from_doc = "dcie" in doc_metrics
        co2_from_doc = ("co2_tonnes" in doc_metrics) or ("total_energy_mwh" in doc_metrics) or ("carbon_factor" in doc_metrics)

        metric_cols = st.columns(3)
        with metric_cols[0]:
            pue_badge = " <span class='kpi-applied'>Applied</span>" if pue_from_doc else ""
            pue_source = "<div class='kpi-source'>Source: Document</div>" if pue_from_doc else ""
            st.markdown(
                f"<div class='metric-card'><h3>"
                f"<svg class='svg-icon' width='14' height='14' viewBox='0 0 24 24' fill='none' "
                f"xmlns='http://www.w3.org/2000/svg'><path d='M13 2L3 14h7l-1 8 10-12h-7l1-8z' "
                f"fill='#cfe0ff'/></svg>PUE{pue_badge}</h3><div class='value'>{pue:.2f}</div>{pue_source}</div>",
                unsafe_allow_html=True,
            )
        with metric_cols[1]:
            dcie_badge = " <span class='kpi-applied'>Applied</span>" if dcie_from_doc else ""
            dcie_source = "<div class='kpi-source'>Source: Document</div>" if dcie_from_doc else ""
            st.markdown(
                f"<div class='metric-card'><h3>"
                f"<svg class='svg-icon' width='14' height='14' viewBox='0 0 24 24' fill='none' "
                f"xmlns='http://www.w3.org/2000/svg'><path d='M4 19h16v2H4z' fill='#cfe0ff'/>"
                f"<path d='M6 17V9h3v8H6zm5 0V5h3v12h-3zm5 0v-6h3v6h-3z' fill='#cfe0ff'/></svg>"
                f"DCiE{dcie_badge}</h3><div class='value'>{dcie:.1f}%</div>{dcie_source}</div>",
                unsafe_allow_html=True,
            )
        with metric_cols[2]:
            co2_badge = " <span class='kpi-applied'>Applied</span>" if co2_from_doc else ""
            co2_source = "<div class='kpi-source'>Source: Document</div>" if co2_from_doc else ""
            st.markdown(
                f"<div class='metric-card'><h3>"
                f"<svg class='svg-icon' width='14' height='14' viewBox='0 0 24 24' fill='none' "
                f"xmlns='http://www.w3.org/2000/svg'><path d='M12 2C7 2 3 6 3 11c0 4 2.6 7.5 6.4 8.7'"
                f" stroke='#cfe0ff' stroke-width='2' fill='none'/>"
                f"<path d='M12 2c5 0 9 4 9 9 0 4-2.6 7.5-6.4 8.7' stroke='#cfe0ff' stroke-width='2' fill='none'/>"
                f"</svg>CO2{co2_badge}</h3><div class='value'>{co2_tonnes:.1f} t/y</div>{co2_source}</div>",
                unsafe_allow_html=True,
            )

        st.markdown(
            f"<div class='section'>Servers: <b>{servers}</b> | CPU Utilization: <b>{cpu_utilization:.1f}%</b>"
            f" | Cooling Setpoint: <b>{cooling_setpoint:.1f} °C</b></div>",
            unsafe_allow_html=True,
        )
        if doc_metrics:
            st.markdown(
                "<div class='subtle'>Document metrics applied to KPIs and simulation.</div>",
                unsafe_allow_html=True,
            )
            if "applied_params" in locals() and applied_params:
                st.markdown("<div class='soft-divider'></div>", unsafe_allow_html=True)
                st.markdown("<div class='subtle'><b>Applied parameters</b></div>", unsafe_allow_html=True)
                for item in applied_params[:6]:
                    st.markdown(f"<div class='subtle'>• {item}</div>", unsafe_allow_html=True)
        st.markdown(
            "<div class='subtle'>Tip: Improve DCiE by reducing non-IT energy overheads.</div>",
            unsafe_allow_html=True,
        )

    with recs_col:
        st.markdown("<div id='recommendations' class='section-title'>AI Recommendations</div>", unsafe_allow_html=True)
        recommendations = build_recommendations(
            cpu_utilization_pct=cpu_utilization,
            cooling_setpoint_c=cooling_setpoint,
            has_aisle_containment=aisle_containment,
            virtualization_level_pct=virtualization_level,
        )
        recs_data = [
            {
                "Action": r.title,
                "Why it helps": r.reason,
                "Estimated Saving (%)": r.estimated_saving_pct,
            }
            for r in recommendations
        ]
        for rec in recs_data:
            st.markdown(
                f"<div class='rec-card'>"
                f"<div class='rec-title'>{action_icon_svg(rec['Action'])}{rec['Action']}</div>"
                f"<div class='subtle'>{rec['Why it helps']}</div>"
                f"<div class='rec-meta'>Estimated Saving: {rec['Estimated Saving (%)']}%</div>"
                f"</div>",
                unsafe_allow_html=True,
            )

    st.markdown("<div id='simulation' class='section-title'>Before / After Simulation</div>", unsafe_allow_html=True)
    actions = [{"estimated_saving_pct": r.estimated_saving_pct} for r in recommendations]
    simulation = simulate_actions(total_energy_mwh, actions)

    sim_df = pd.DataFrame(
        [
            {"Scenario": "Baseline", "Energy (MWh/year)": simulation["initial_energy_mwh"]},
            {
                "Scenario": "After Actions",
                "Energy (MWh/year)": simulation["remaining_energy_mwh"],
            },
        ]
    )
    st.bar_chart(sim_df.set_index("Scenario"))

    st.markdown(
        f"<div class='section'>Estimated total savings: <b>{simulation['total_savings_mwh']:.1f} MWh/year</b>"
        f" ({simulation['total_savings_pct']:.1f}%)</div>",
        unsafe_allow_html=True,
    )
    if simulation["total_savings_pct"] >= 25:
        st.success("Target -25% CO2 is achievable with these actions.")
    else:
        st.info("Target -25% CO2 not reached. Adjust the action set.")
    doc_texts = st.session_state.get("doc_texts", [])
    if doc_texts:
        inputs = extract_workload_inputs("\n".join(doc_texts))
        results = compute_workload_audit(inputs)
        if "error" not in results:
            n_value = inputs.get("n_inferences")
            n_display = f"{int(n_value):,}" if n_value else "n/a"
            decision = summarize_workload_decision(inputs, results)
            recommendation = workload_recommendation(inputs, results)
            flags = workload_decision_flags(results)
            goal = st.session_state.get("business_goal", "Not specified")
            goal_line = business_goal_recommendation(goal, results)
            rec_html = "<br>".join([line.replace("- ", "• ") for line in recommendation.splitlines()])
            st.markdown("<div class='section-title'>AI Workload Audit</div>", unsafe_allow_html=True)
            st.markdown(
                f"""
                <div class='section'>
                N = <b>{n_display}</b> inferences (total inferences in the scenario)<br>
                GPU total time (h): <b>{results['gpu_time_h']:.2f}</b><br>
                Edge total time (h): <b>{results['edge_time_h']:.2f}</b><br>
                GPU total energy (kWh): <b>{results['gpu_energy_kwh']:.3f}</b><br>
                Edge total energy (kWh): <b>{results['edge_energy_kwh']:.3f}</b><br>
                GPU energy per inference (Wh): <b>{results['gpu_energy_per_inf_wh']:.6f}</b><br>
                Edge energy per inference (Wh): <b>{results['edge_energy_per_inf_wh']:.6f}</b><br>
                GPU cost (€): <b>{results['gpu_cost']:.2f}</b><br>
                Edge cost (€): <b>{results['edge_cost']:.2f}</b>
                <br><br><b>Decision:</b><br>{decision.replace("\n", "<br>")}<br><br>
                <b>Business goal:</b><br>{goal_line}<br><br>
                <b>Key indicators:</b><br>
                <span class='badge badge-solid'>Fastest: {flags['fastest']}</span>
                <span class='badge badge-solid'>Lowest energy: {flags['lowest_energy']}</span>
                <span class='badge badge-solid'>Lowest cost: {flags['lowest_cost']}</span>
                <br><br><b>Recommendation:</b><br>{rec_html}
                </div>
                """,
                unsafe_allow_html=True,
            )
    st.markdown(
        "<div class='footer'>GreenDC Audit Platform • Responsible by design • © GreenAI Systems</div>",
        unsafe_allow_html=True,
    )

    if "assistant_visible" not in st.session_state:
        st.session_state.assistant_visible = True
    if st.session_state.assistant_visible:
        st.markdown('<div id="assistant" class="section-title">GreenDC Audit AI</div>', unsafe_allow_html=True)
        st.markdown("<div class='soft-divider'></div>", unsafe_allow_html=True)
        with st.form("assistant_form", clear_on_submit=False):
            question = st.text_area(
                "Ask a question about your audit (e.g., how to reach -25% CO2?)",
                height=90,
                key="assistant_question",
            )
            submitted = st.form_submit_button("Ask Knowledge Assistant")
        if submitted:
            context = {
                "it_energy_mwh": it_energy_mwh,
                "total_energy_mwh": total_energy_mwh,
                "carbon_factor": carbon_factor,
                "pue": pue,
                "dcie": dcie,
                "co2_tonnes": co2_tonnes,
                "cpu_utilization": cpu_utilization,
                "cooling_setpoint": cooling_setpoint,
                "aisle_containment": aisle_containment,
                "virtualization_level": virtualization_level,
                "recommendations": recs_data,
                "cooling_ratio": 100.0 * (1 - (1 / pue)) if pue > 0 else 0.0,
                "doc_summaries": st.session_state.get("doc_summaries", []),
                "doc_metrics": st.session_state.get("doc_metrics", {}),
            }
            rules = load_rules()
            reply = ai_assistant_reply(question, context)
            is_doc_query = is_document_question(question) or is_long_paste(question)
            if (is_audit_question(question) or context.get("doc_metrics")) and not is_doc_query:
                rules_block = rule_based_plan(context, rules)
                reply += "\n\n" + rules_block
            docs_used = "YES" if context.get("doc_summaries") else "NO"
            reply += f"\n\nDocs used: {docs_used}"
            if simulate_web and not is_doc_query:
                reply += "\n\nSimulated web search: This is a mock summary based on best practices."
            st.session_state.assistant_reply = reply
        if "assistant_reply" in st.session_state:
            st.markdown(f"<div class='section'>{st.session_state.assistant_reply}</div>", unsafe_allow_html=True)
        st.caption("Offline assistant. Uses rules + local/HF knowledge only (no web scraping).")

    st.markdown("<div class='section-title'>Document QA</div>", unsafe_allow_html=True)
    st.markdown("<div class='soft-divider'></div>", unsafe_allow_html=True)
    business_goal = st.selectbox(
        "Business goal (optional)",
        ["Not specified", "Minimum cost", "Minimum energy per inference", "Minimum latency"],
        index=0,
    )
    st.session_state["business_goal"] = business_goal
    doc_question = st.text_area(
        "Ask a question about uploaded documents (e.g., extract latency or cost KPIs).",
        height=90,
        key="doc_question",
    )
    if st.button("Analyze Documents"):
        docs = st.session_state.get("doc_texts", [])
        combined = "\n".join(docs) + "\n" + doc_question
        if not docs and not doc_question.strip():
            st.session_state["doc_reply"] = "No document uploaded yet."
        else:
            inputs = extract_workload_inputs(combined)
            results = compute_workload_audit(inputs)
            if "error" in results:
                if any(inputs.get(k) is not None for k in ["n_inferences", "gpu_power_w", "edge_power_w", "gpu_latency_ms", "edge_latency_ms"]):
                    st.session_state["doc_reply"] = (
                        "Document KPIs detected but incomplete. "
                        + results["error"]
                        + "\nExtracted values: "
                        + ", ".join(f"{k}={v}" for k, v in inputs.items())
                    )
                else:
                    st.session_state["doc_reply"] = (
                        "This document is a Green IT audit scenario (not an AI workload benchmark). "
                        "Use the main KPIs and Applied parameters above."
                    )
            else:
                gpu_cost_line = (
                    f"GPU cost (€): {results['gpu_cost']:.2f}"
                    if results["gpu_cost"] is not None
                    else "GPU cost (€): n/a"
                )
                edge_cost_line = (
                    f"Edge cost (€): {results['edge_cost']:.2f}"
                    if results["edge_cost"] is not None
                    else "Edge cost (€): n/a"
                )
                decision = summarize_workload_decision(inputs, results)
                recommendation = workload_recommendation(inputs, results)
                goal_line = business_goal_recommendation(business_goal, results)
                n_value = inputs.get("n_inferences")
                n_display = f"{int(n_value):,}" if n_value else "n/a"
                rec_html = "<br>".join([line.replace("- ", "• ") for line in recommendation.splitlines()])
                st.session_state["doc_reply"] = (
                    "<div class='section'>"
                    "<b>AI Workload Audit (from documents)</b><br>"
                    f"N = <b>{n_display}</b> inferences (total inferences in the scenario)<br>"
                    f"GPU total time (h): <b>{results['gpu_time_h']:.2f}</b><br>"
                    f"Edge total time (h): <b>{results['edge_time_h']:.2f}</b><br>"
                    f"GPU total energy (kWh): <b>{results['gpu_energy_kwh']:.3f}</b><br>"
                    f"Edge total energy (kWh): <b>{results['edge_energy_kwh']:.3f}</b><br>"
                    f"GPU energy per inference (Wh): <b>{results['gpu_energy_per_inf_wh']:.6f}</b><br>"
                    f"Edge energy per inference (Wh): <b>{results['edge_energy_per_inf_wh']:.6f}</b><br>"
                    f"{gpu_cost_line}<br>"
                    f"{edge_cost_line}"
                    "<div class='soft-divider'></div>"
                    "<b>Decision</b><br>"
                    f"{decision.replace('\n', '<br>')}"
                    "<div class='soft-divider'></div>"
                    f"<b>Business goal</b><br>{goal_line}"
                    "<div class='soft-divider'></div>"
                    "<b>Recommendation</b><br>"
                    f"{rec_html}"
                    "</div>"
                )
    if st.session_state.get("doc_reply"):
        st.markdown(st.session_state["doc_reply"], unsafe_allow_html=True)

if page == "About":
    st.markdown("<div id='about' class='section-title'>About the Platform</div>", unsafe_allow_html=True)
    st.markdown(
        """
        <div class="glass">
            <b>Mission:</b> Provide an audit-ready, measurable path to cut data center CO2 by 25%.
            <br><br>
            <b>Principles:</b> Digital sobriety, proportional computing, and continuous measurement.
            <br><br>
            <b>Scope:</b> PUE, DCiE, CO2, recommendations, and scenario validation.
        </div>
        """,
        unsafe_allow_html=True,
    )
    st.markdown("")
    st.markdown(
        """
        <div class="glass">
            <b>How to use:</b>
            <ol>
                <li>Enter data center inputs in the sidebar.</li>
                <li>Review metrics and AI recommendations.</li>
                <li>Validate the -25% target with the simulation panel.</li>
            </ol>
        </div>
        """,
        unsafe_allow_html=True,
    )
    st.markdown("")
    st.markdown(
        """
        <div class="glass">
            <b>Modules:</b> frontend, energy_metrics, ai_recommendation, simulation, case_study.
            <br><br>
            <span class="badge badge-solid">Green IT</span>
            <span class="badge badge-solid">Green Coding</span>
            <span class="badge badge-solid">Proportional Computing</span>
        </div>
        """,
        unsafe_allow_html=True,
    )
    st.markdown("<div id='team' class='section-title'>Team</div>", unsafe_allow_html=True)
    st.markdown(
        """
        <div class="glass">
            <b>GreenAI Systems</b><br><br>
            Gemima Ondele Pourou • Platform Architect & Frontend/Integration<br>
            Joseph Fabrice Tsapfack • AI & Recommendation Engine<br>
            Mike-Brady • Simulation & Validation<br>
            Nandaa • Data & Case Study<br>
            Pierre Joel • Documentation & QA
        </div>
        """,
        unsafe_allow_html=True,
    )
